import os
import re
import uuid
import time
import logging
import tempfile
import requests
from datetime import datetime
from typing import Dict, Any, Optional, List
from bs4 import BeautifulSoup

import streamlit as st
from pypdf import PdfReader
import docx

from config import MAX_CONTENT_SIZE, REQUEST_TIMEOUT, WEB_BATCH_SIZE
from utils import estimate_text_density, split_into_chunks, MemoryManager
from utils.sitemap_utils import get_sitemap_urls

def process_text(text_content: str, source_metadata: Optional[Dict] = None, instance=None):
    """Process raw text and add to RAG instance."""
    if not instance:
        instance = st.session_state.rag_instances.get(st.session_state.current_rag_instance)
        
    if not instance:
        st.error("No RAG instance selected for processing text.")
        return 0

    if len(text_content) > MAX_CONTENT_SIZE:
        logging.warning(f"Text size ({len(text_content)} bytes) exceeds limit. Truncating.")
        text_content = text_content[:MAX_CONTENT_SIZE]

    chunk_size = estimate_text_density(text_content)
    chunks = split_into_chunks(text_content, chunk_size=chunk_size)
    if not chunks: return 0

    metadatas = [(source_metadata.copy() if source_metadata else {}) for _ in chunks]
    
    ids = instance.add_texts(chunks, metadatas)
    return len(ids)


def process_pdf(pdf_file_path: str, filename: str, instance=None):
    """Process PDF document and add to RAG instance."""
    if not instance:
        instance = st.session_state.rag_instances.get(st.session_state.current_rag_instance)
        
    if not instance:
        st.error("No RAG instance selected for processing PDF.")
        return 0
    
    all_chunks = []
    all_metadatas = []

    try:
        reader = PdfReader(pdf_file_path)
        source_metadata_base = {
            "source_type": "pdf",
            "filename": filename,
            "date_added": datetime.now().isoformat(),
            "total_pages": len(reader.pages)
        }
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        num_pages = len(reader.pages)

        for i, page in enumerate(reader.pages):
            progress = (i + 1) / num_pages
            progress_bar.progress(progress)
            status_text.text(f"Processing PDF page {i+1} of {num_pages}")
            try:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    page_metadata = source_metadata_base.copy()
                    page_metadata["page"] = i + 1
                    
                    chunk_size = estimate_text_density(page_text)
                    page_chunks_content = split_into_chunks(page_text, chunk_size=chunk_size)
                    
                    for chunk_content in page_chunks_content:
                        all_chunks.append(chunk_content)
                        all_metadatas.append(page_metadata.copy())
            except Exception as e:
                logging.error(f"Error processing PDF page {i+1} of {filename}: {str(e)}", exc_info=True)
        
        if all_chunks:
            instance.add_texts(all_chunks, all_metadatas)
            instance.add_document({
                "type": "pdf", "filename": filename, "chunks": len(all_chunks), 
                "size": os.path.getsize(pdf_file_path) if os.path.exists(pdf_file_path) else 0,
                "date_added": datetime.now().isoformat()
            })

        progress_bar.empty()
        status_text.empty()
        return len(all_chunks)
    except Exception as e:
        logging.error(f"Error processing PDF file {filename}: {str(e)}", exc_info=True)
        st.error(f"Error processing PDF '{filename}': {str(e)}")
        if 'progress_bar' in locals(): progress_bar.empty()
        if 'status_text' in locals(): status_text.empty()
        return 0


def process_docx(docx_file_path: str, filename: str, instance=None):
    """Process DOCX document and add to RAG instance."""
    if not instance:
        instance = st.session_state.rag_instances.get(st.session_state.current_rag_instance)
        
    if not instance:
        st.error("No RAG instance selected for processing DOCX.")
        return 0

    all_chunks = []
    all_metadatas = []

    try:
        doc = docx.Document(docx_file_path)
        source_metadata_base = {
            "source_type": "docx",
            "filename": filename,
            "date_added": datetime.now().isoformat(),
            "total_paragraphs": len(doc.paragraphs)
        }

        progress_bar = st.progress(0)
        status_text = st.empty()
        total_paragraphs = len(doc.paragraphs)
        
        # Process paragraph by paragraph to build up text, then chunk that text
        full_text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        if full_text_content:
            status_text.text(f"Chunking content from '{filename}'...")
            chunk_size = estimate_text_density(full_text_content)
            doc_chunks_content = split_into_chunks(full_text_content, chunk_size=chunk_size)
            
            for i, chunk_content in enumerate(doc_chunks_content):
                progress = (i + 1) / len(doc_chunks_content) if doc_chunks_content else 1
                progress_bar.progress(progress)
                chunk_metadata = source_metadata_base.copy()
                chunk_metadata["chunk_index"] = i + 1
                all_chunks.append(chunk_content)
                all_metadatas.append(chunk_metadata)

        if all_chunks:
            instance.add_texts(all_chunks, all_metadatas)
            instance.add_document({
                "type": "docx", "filename": filename, "chunks": len(all_chunks),
                "size": os.path.getsize(docx_file_path) if os.path.exists(docx_file_path) else 0,
                "date_added": datetime.now().isoformat()
            })

        progress_bar.empty()
        status_text.empty()
        return len(all_chunks)
    except Exception as e:
        logging.error(f"Error processing DOCX file {filename}: {str(e)}", exc_info=True)
        st.error(f"Error processing DOCX '{filename}': {str(e)}")
        if 'progress_bar' in locals(): progress_bar.empty()
        if 'status_text' in locals(): status_text.empty()
        return 0


def process_url(url: str, max_size=MAX_CONTENT_SIZE, use_sitemap=True, instance=None):
    """
    Enhanced URL processing with optional sitemap support.
    If use_sitemap is True, will attempt to find and process URLs from XML sitemaps.
    """
    if not instance:
        instance = st.session_state.rag_instances.get(st.session_state.current_rag_instance)
        
    if not instance:
        st.error("No RAG instance selected for processing URL.")
        return 0

    all_url_chunks = []
    all_url_metadatas = []
    processed_urls = set()  # Track processed URLs to avoid duplicates
    total_chunks = 0
    
    # Function to process a single URL
    def process_single_url(url, is_main_url=False):
        nonlocal total_chunks
        try:
            if url in processed_urls:
                return 0  # Skip already processed URLs
                
            processed_urls.add(url)
            
            headers = {"User-Agent": "Mozilla/5.0 (compatible; StreamlitChatbot/1.0; +http://localhost)"}
            content_type = ""
            content_length = 0

            try:
                with requests.head(url, timeout=REQUEST_TIMEOUT, headers=headers, allow_redirects=True) as head_response:
                    head_response.raise_for_status()
                    content_type = head_response.headers.get("content-type", "")
                    content_length = int(head_response.headers.get("content-length", 0))
                    
                    if not content_type.startswith("text/html") and not is_main_url:
                        logging.info(f"Skipping non-HTML URL: {url} ({content_type})")
                        return 0
                    
                    if content_length > max_size * 1.5:  # Add some margin
                         if is_main_url:
                             st.warning(f"URL content is large ({content_length/1024/1024:.2f}MB). Processing may take time or be partial.")
                         else:
                             logging.info(f"Skipping large URL: {url} ({content_length/1024/1024:.2f}MB)")
                             return 0
            except requests.exceptions.RequestException as e:
                if is_main_url:
                    logging.warning(f"HEAD request for URL {url} failed: {e}. Proceeding with GET.")
                    st.info("Could not determine URL size/type in advance. Processing with care.")
                else:
                    logging.warning(f"Skipping URL due to HEAD error: {url}")
                    return 0

            source_metadata_base = {
                "source_type": "url", "url": url, "date_added": datetime.now().isoformat(), "title": url
            }
            
            # Only display progress for main URL
            progress_bar = st.progress(0) if is_main_url else None
            status_text = st.empty() if is_main_url else None
            if is_main_url:
                status_text.text(f"Downloading content from {url}...")

            with requests.get(url, timeout=REQUEST_TIMEOUT, headers=headers, stream=True, allow_redirects=True) as response:
                response.raise_for_status()
                
                # Update content_type and content_length if not available from HEAD
                if not content_type: content_type = response.headers.get("content-type", "")
                if content_length == 0: content_length = int(response.headers.get("content-length", 0))

                content_buffer = b""
                total_size_processed = 0
                
                for chunk_bytes in response.iter_content(chunk_size=WEB_BATCH_SIZE, decode_unicode=False):
                    if not chunk_bytes: continue
                    content_buffer += chunk_bytes
                    total_size_processed += len(chunk_bytes)

                    if is_main_url:
                        if content_length > 0:
                            progress = min(total_size_processed / content_length, 1.0)
                            progress_bar.progress(progress)
                            status_text.text(f"Downloaded {total_size_processed/1024:.1f}KB of {content_length/1024:.1f}KB")
                        else:
                            status_text.text(f"Downloaded {total_size_processed/1024:.1f}KB")
                    
                    if total_size_processed >= max_size:
                        if is_main_url:
                            st.warning(f"Reached processing limit of {max_size/1024/1024:.1f}MB for URL. Content may be truncated.")
                        break
                
                # Decode the entire buffer (up to max_size)
                try:
                    decoded_content = content_buffer.decode('utf-8', errors='replace')
                except UnicodeDecodeError:
                    decoded_content = content_buffer.decode('latin-1', errors='replace')

                if is_main_url:
                    status_text.text("Parsing HTML content...")
                
                text_to_chunk = decoded_content
                if "text/html" in content_type.lower():
                    soup = BeautifulSoup(decoded_content, 'html.parser')
                    if soup.title and soup.title.string:
                        source_metadata_base["title"] = soup.title.string.strip()
                    
                    for element in soup(["script", "style", "header", "footer", "nav", "aside", "form", "button", "input", "select", "textarea"]):
                        element.extract()
                    
                    main_content_area = soup.find('main') or soup.find('article') or soup.body
                    text_to_chunk = main_content_area.get_text(separator="\n\n", strip=True) if main_content_area else soup.get_text(separator="\n\n", strip=True)

                if text_to_chunk.strip():
                    if is_main_url:
                        status_text.text(f"Chunking content from '{source_metadata_base['title']}'...")
                    chunk_size_est = estimate_text_density(text_to_chunk)
                    url_page_chunks = split_into_chunks(text_to_chunk, chunk_size=chunk_size_est)
                    
                    url_chunks = []
                    url_metadatas = []
                    
                    for i, chunk_content in enumerate(url_page_chunks):
                        chunk_meta = source_metadata_base.copy()
                        chunk_meta["chunk_index"] = i + 1
                        url_chunks.append(chunk_content)
                        url_metadatas.append(chunk_meta)
                    
                    # Add to the global tracking
                    all_url_chunks.extend(url_chunks)
                    all_url_metadatas.extend(url_metadatas)
                    
                    # Track statistics
                    total_chunks += len(url_chunks)
                    
                    # If this isn't the main URL, directly add to the database
                    if not is_main_url and url_chunks:
                        instance.add_texts(url_chunks, url_metadatas)
                        document_info = {
                            "type": "url",
                            "url": url,
                            "title": source_metadata_base["title"],
                            "chunks": len(url_chunks),
                            "size": total_size_processed,
                            "date_added": datetime.now().isoformat(),
                            "from_sitemap": True
                        }
                        instance.add_document(document_info)
            
            if is_main_url:
                progress_bar.empty()
                status_text.empty()
            
            return len(url_chunks)

        except requests.exceptions.RequestException as e:
            if is_main_url:
                logging.error(f"Error processing URL {url} (RequestException): {str(e)}", exc_info=True)
                st.error(f"Could not fetch URL {url}: {str(e)}")
            return 0
        except Exception as e:
            if is_main_url:
                logging.error(f"Error processing URL {url} (General): {str(e)}", exc_info=True)
                st.error(f"Error processing URL {url}: {str(e)}")
            return 0
        finally:
            if is_main_url:
                if 'progress_bar' in locals() and progress_bar: progress_bar.empty()
                if 'status_text' in locals() and status_text: status_text.empty()
                
    # First, process the main URL
    main_url_chunks = process_single_url(url, is_main_url=True)
    
    # Process sitemap if requested
    if use_sitemap and main_url_chunks > 0:
        sitemap_progress = st.progress(0)
        sitemap_status = st.empty()
        sitemap_status.text("Looking for sitemap...")
        
        try:
            sitemap_urls = get_sitemap_urls(url)
            
            if sitemap_urls:
                sitemap_status.text(f"Found {len(sitemap_urls)} URLs in sitemap. Processing...")
                
                # Process each URL from the sitemap
                for i, sitemap_url in enumerate(sitemap_urls):
                    progress = (i + 1) / len(sitemap_urls)
                    sitemap_progress.progress(progress)
                    sitemap_status.text(f"Processing sitemap URL {i+1}/{len(sitemap_urls)}: {sitemap_url[:30]}...")
                    process_single_url(sitemap_url)
                
                sitemap_status.text(f"Completed processing {len(sitemap_urls)} URLs from sitemap.")
            else:
                sitemap_status.text("No sitemap found or no URLs extracted from sitemap.")
                
        except Exception as e:
            logging.error(f"Error processing sitemap for {url}: {str(e)}", exc_info=True)
            sitemap_status.text(f"Error processing sitemap: {str(e)}")
        finally:
            sitemap_progress.empty()
            sitemap_status.empty()
    
    # Add the main page chunks if any
    if all_url_chunks:
        instance.add_texts(all_url_chunks, all_url_metadatas)
        instance.add_document({
            "type": "url", 
            "url": url, 
            "title": all_url_metadatas[0].get("title", url) if all_url_metadatas else url, 
            "chunks": len(all_url_chunks),
            "total_urls_processed": len(processed_urls),
            "size": sum(1 for _ in all_url_chunks),  # Simple count for size
            "date_added": datetime.now().isoformat()
        })

    return total_chunks


def process_uploaded_file(uploaded_file, instance=None):
    """Process an uploaded file based on its type."""
    if not uploaded_file: return 0
    
    if not instance:
        instance = st.session_state.rag_instances.get(st.session_state.current_rag_instance)
        
    if not instance:
        st.error("No RAG instance selected for processing file.")
        return 0
        
    try:
        # Use a temporary file to handle various file types robustly
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_f:
            temp_f.write(uploaded_file.getvalue())
            temp_file_path = temp_f.name
        
        chunk_count = 0
        try:
            if uploaded_file.name.lower().endswith(".pdf"):
                chunk_count = process_pdf(temp_file_path, uploaded_file.name, instance)
            elif uploaded_file.name.lower().endswith(".docx"):
                chunk_count = process_docx(temp_file_path, uploaded_file.name, instance)
            elif uploaded_file.name.lower().endswith((".txt", ".md", ".py", ".js", ".html", ".css")):  # Expand to common text types
                with open(temp_file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
                chunk_count = process_text(text, {
                    "source_type": "text_file", 
                    "filename": uploaded_file.name,
                    "date_added": datetime.now().isoformat()
                }, instance)
                if chunk_count > 0:  # For text files, add document info manually
                    instance.add_document({
                        "type": "text_file", "filename": uploaded_file.name, "chunks": chunk_count,
                        "size": uploaded_file.size, "date_added": datetime.now().isoformat()
                    })
            else:
                st.warning(f"Unsupported file type: {uploaded_file.name}. Supported: PDF, DOCX, TXT, MD, code files.")

            return chunk_count
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)  # Clean up temp file
    except Exception as e:
        logging.error(f"Error processing uploaded file {uploaded_file.name}: {str(e)}", exc_info=True)
        st.error(f"Error processing file '{uploaded_file.name}': {str(e)}")
        return 0